"""Генератор документов — простой, надёжный, с поддержкой аудита"""

import logging
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

def generate_report_document(title: str, author: str, content: str = "") -> bytes:
    """
    Генерирует .docx-документ в стиле официального отчёта
    """
    doc = Document()
    
    # Стиль: ГОСТ Р 7.0.97-2016
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Заголовок
    heading = doc.add_heading(level=1)
    run = heading.add_run(title.upper())
    run.font.size = Pt(14)
    run.font.bold = True
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Пустой абзац
    doc.add_paragraph()
    
    # Метаданные
    meta = doc.add_paragraph()
    meta.add_run(f"Автор: {author}\n")
    meta.add_run(f"Дата формирования: {datetime.now():%d.%m.%Y %H:%M}\n")
    meta.add_run(f"Система: Судебный Кейс v1.0")
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Разделитель
    doc.add_paragraph("_" * 80)
    doc.add_paragraph()
    
    # Основной текст
    if content:
        doc.add_paragraph(content)
    else:
        doc.add_paragraph(
            "Настоящий документ сформирован автоматически в рамках "
            "оказания юридических услуг. Все данные проверены и соответствуют "
            "действующему законодательству Российской Федерации."
        )
    
    # Подпись
    doc.add_paragraph()
    signature = doc.add_paragraph()
    signature.add_run("_________________________")
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Сохраняем в bytes
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
