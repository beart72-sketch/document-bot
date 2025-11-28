#!/usr/bin/env python3
"""
Создает правильный шаблон для docxtpl
"""

from docx import Document
from docx.shared import Inches

def create_contract_template():
    # Создаем новый документ
    doc = Document()
    
    # Добавляем заголовок
    title = doc.add_heading('ДОГОВОР № {{ contract_number }}', 0)
    
    # Добавляем место и дату
    doc.add_paragraph('г. {{ city }}                                 «___»__________ 20___ г.')
    doc.add_paragraph('')
    
    # Стороны договора
    p = doc.add_paragraph()
    p.add_run('{{ client_name }}').bold = True
    p.add_run(', именуем__ в дальнейшем «Заказчик», в лице ')
    p.add_run('{{ client_representative }}').bold = True
    p.add_run(', действующ___ на основании ')
    p.add_run('{{ client_basis }}').bold = True
    p.add_run(', с одной стороны, и ')
    
    p = doc.add_paragraph()
    p.add_run('{{ executor_name }}').bold = True  
    p.add_run(', именуем__ в дальнейшем «Исполнитель», в лице ')
    p.add_run('{{ executor_representative }}').bold = True
    p.add_run(', действующ___ на основании ')
    p.add_run('{{ executor_basis }}').bold = True
    p.add_run(', с другой стороны, заключили настоящий Договор о нижеследующем:')
    
    doc.add_paragraph('')
    
    # Предмет договора
    doc.add_heading('1. ПРЕДМЕТ ДОГОВОРА', level=1)
    doc.add_paragraph('{{ contract_subject }}')
    doc.add_paragraph('')
    
    # Стоимость
    doc.add_heading('2. СТОИМОСТЬ УСЛУГ И ПОРЯДОК РАСЧЕТОВ', level=1)
    doc.add_paragraph('{{ payment_terms }}')
    doc.add_paragraph('')
    
    # Срок действия
    doc.add_heading('3. СРОК ДЕЙСТВИЯ ДОГОВОРА', level=1) 
    doc.add_paragraph('{{ contract_term }}')
    doc.add_paragraph('')
    
    # Реквизиты
    doc.add_heading('4. ЮРИДИЧЕСКИЕ АДРЕСА И РЕКВИЗИТЫ СТОРОН', level=1)
    doc.add_paragraph('Заказчик: {{ client_details }}')
    doc.add_paragraph('')
    doc.add_paragraph('Исполнитель: {{ executor_details }}')
    doc.add_paragraph('')
    
    # Подписи
    doc.add_paragraph('ПОДПИСИ СТОРОН:')
    doc.add_paragraph('')
    doc.add_paragraph('Заказчик: _________________________/{{ client_signature }}/')
    doc.add_paragraph('')
    doc.add_paragraph('Исполнитель: _______________________/{{ executor_signature }}/')
    doc.add_paragraph('')
    
    # Дата создания
    doc.add_paragraph(f'Шаблон создан: {__import__("datetime").datetime.now().strftime("%d.%m.%Y %H:%M")}')
    
    # Сохраняем
    doc.save('templates/contract_template.docx')
    print("✅ Шаблон создан: templates/contract_template.docx")

if __name__ == "__main__":
    create_contract_template()
